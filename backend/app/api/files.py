"""
Endpoint de upload e indexacao de documentos (POST /api/files/upload).
Extrai texto de PDF, DOCX ou imagens (nome/metadados) e envia para
indexacao vetorial (RAG), para que o conteudo passe a servir de contexto
nas conversas do usuario.
"""
import os
import uuid

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from app.auth.deps import get_current_user
from app.config import get_settings
from app.database import AsyncSessionLocal, get_db
from app.models import Document, User
from app.vector_database.store import index_document

router = APIRouter(prefix="/api/files", tags=["files"])
settings = get_settings()

ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/png": "image",
    "image/jpeg": "image",
}


def _extract_text(path: str, kind: str) -> str:
    if kind == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if kind == "docx":
        import docx
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    if kind == "txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    if kind == "image":
        # Imagens nao tem texto extraivel sem OCR/visao; guardamos apenas o nome
        # como referencia de contexto (a IA pode receber a imagem diretamente
        # no futuro via input multimodal, dependendo do modelo configurado).
        return f"[Imagem enviada pelo usuario: {os.path.basename(path)}]"
    return ""


async def _process_upload(document_id: str, path: str, kind: str):
    async with AsyncSessionLocal() as db:
        from sqlalchemy import select
        result = await db.execute(select(Document).where(Document.id == document_id))
        document = result.scalar_one_or_none()
        if not document:
            return
        try:
            text = _extract_text(path, kind)
            await index_document(db, document, text)
        except Exception:
            document.status = "error"
            await db.commit()


@router.post("/upload")
async def upload_file(
    file: UploadFile,
    background_tasks: BackgroundTasks,
    conversation_id: str | None = None,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail="Tipo de arquivo nao suportado")

    contents = await file.read()
    size_mb = len(contents) / (1024 * 1024)
    if size_mb > settings.MAX_UPLOAD_MB:
        raise HTTPException(status_code=400, detail=f"Arquivo excede o limite de {settings.MAX_UPLOAD_MB}MB")

    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    stored_name = f"{uuid.uuid4()}_{file.filename}"
    path = os.path.join(settings.UPLOAD_DIR, stored_name)
    with open(path, "wb") as f:
        f.write(contents)

    document = Document(
        user_id=user.id,
        conversation_id=conversation_id,
        filename=file.filename,
        content_type=file.content_type,
        size_bytes=len(contents),
        status="processing",
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    kind = ALLOWED_TYPES[file.content_type]
    background_tasks.add_task(_process_upload, document.id, path, kind)

    return {"id": document.id, "filename": document.filename, "status": document.status}


@router.get("/{document_id}/status")
async def upload_status(document_id: str, user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    from sqlalchemy import select
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.user_id == user.id)
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=404, detail="Documento nao encontrado")
    return {"id": document.id, "status": document.status, "filename": document.filename}
